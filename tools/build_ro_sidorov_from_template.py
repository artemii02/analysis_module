from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph


BASE = Path("C:/Users/super")
REF_PREFIX = "\u0420\u041e \u0411\u0435\u0440\u0435\u0437\u0443\u0442"
REFERENCE = next(
    p for p in (BASE / "Downloads/Telegram Desktop").glob("*.docx") if p.name.startswith(REF_PREFIX)
)
TARGET = (
    BASE
    / "Desktop"
    / "\u0414\u0438\u043f\u043b\u043e\u043c"
    / "\u0444\u0430\u0439\u043d\u0430\u043b\u0438 \u0434\u043e\u043a\u0438"
    / "\u0420\u041e_\u0421\u0438\u0434\u043e\u0440\u043e\u0432.docx"
)
BACKUP = TARGET.with_name("\u0420\u041e_\u0421\u0438\u0434\u043e\u0440\u043e\u0432_backup_before_ber_format.docx")

FONT = "Times New Roman"

MODULE_TITLE_LINES = [
    "\u0412\u0415\u0411-\u041f\u0420\u0418\u041b\u041e\u0416\u0415\u041d\u0418\u0415 \u0414\u041b\u042f \u0418\u041c\u0418\u0422\u0410\u0426\u0418\u0418 \u0421\u041e\u0411\u0415\u0421\u0415\u0414\u041e\u0412\u0410\u041d\u0418\u0419",
    "\u0421 \u0418\u041d\u0422\u0415\u041b\u041b\u0415\u041a\u0422\u0423\u0410\u041b\u042c\u041d\u041e\u0419 \u041e\u0411\u0420\u0410\u0422\u041d\u041e\u0419 \u0421\u0412\u042f\u0417\u042c\u042e:",
    "\u041c\u041e\u0414\u0423\u041b\u042c \u0410\u041d\u0410\u041b\u0418\u0417\u0410 \u041e\u0422\u0412\u0415\u0422\u041e\u0412 \u041f\u041e\u041b\u042c\u0417\u041e\u0412\u0410\u0422\u0415\u041b\u042f",
    "\u0421 \u0418\u0421\u041f\u041e\u041b\u042c\u0417\u041e\u0412\u0410\u041d\u0418\u0415\u041c \u0411\u041e\u041b\u042c\u0428\u041e\u0419 \u042f\u0417\u042b\u041a\u041e\u0412\u041e\u0419 \u041c\u041e\u0414\u0415\u041b\u0418",
]


def set_run_font(run, size: int = 14, bold: bool | None = None) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def normalize_paragraph(paragraph: Paragraph, *, heading: bool = False, center: bool = False) -> None:
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif heading:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        set_run_font(run, size=14, bold=heading)


def set_text(paragraph: Paragraph, text: str, *, heading: bool = False, center: bool = False) -> None:
    paragraph.text = text
    normalize_paragraph(paragraph, heading=heading, center=center)


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(block: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    block._p.addnext(new_p)
    paragraph = Paragraph(new_p, block._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    normalize_paragraph(paragraph, heading=bool(style and style.startswith("Heading")))
    return paragraph


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def replace_title_pages(doc: Document) -> None:
    paragraphs = doc.paragraphs
    # First title block.
    for idx, line in zip([27, 28, 29, 30], MODULE_TITLE_LINES, strict=True):
        set_text(paragraphs[idx], line, center=True)
    insert_paragraph_after(paragraphs[30], "\u0420\u0443\u043a\u043e\u0432\u043e\u0434\u0441\u0442\u0432\u043e \u043e\u043f\u0435\u0440\u0430\u0442\u043e\u0440\u0430")
    # Remove duplicated old title paragraph from the template if present.
    if paragraphs[31].text.strip() == "\u0420\u0443\u043a\u043e\u0432\u043e\u0434\u0441\u0442\u0432\u043e \u043e\u043f\u0435\u0440\u0430\u0442\u043e\u0440\u0430":
        pass

    # Second title block. After insertion indices can shift, so search by existing code line.
    title_code = find_paragraph(doc, "RU.17701729.05.15-01 34 01-1-\u041b\u0423")
    approved_heading = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("\u0423\u0422\u0412\u0415\u0420\u0416\u0414"):
            approved_heading = paragraph
            break
    if approved_heading is not None:
        current = approved_heading
        # Replace the four visible title lines after the approval block.
        seen = 0
        for paragraph in doc.paragraphs:
            if paragraph._element.getparent() is None:
                continue
            if seen < 4 and paragraph.text.strip() in {
                "\u0412\u0415\u0411-\u041f\u0420\u0418\u041b\u041e\u0416\u0415\u041d\u0418\u0415 \u0414\u041b\u042f \u0418\u041c\u0418\u0422\u0410\u0426\u0418\u0418 \u0421\u041e\u0411\u0415\u0421\u0415\u0414\u041e\u0412\u0410\u041d\u0418\u0419",
                "\u0421 \u0418\u041d\u0422\u0415\u041b\u041b\u0415\u041a\u0422\u0423\u0410\u041b\u042c\u041d\u041e\u0419 \u041e\u0411\u0420\u0410\u0422\u041d\u041e\u0419 \u0421\u0412\u042f\u0417\u042c\u042e:",
                "\u041a\u041b\u0418\u0415\u041d\u0422\u0421\u041a\u0410\u042f \u0418 \u0421\u0415\u0420\u0412\u0415\u0420\u041d\u0410\u042f \u0427\u0410\u0421\u0422\u042c",
            }:
                set_text(paragraph, MODULE_TITLE_LINES[min(seen, len(MODULE_TITLE_LINES) - 1)], center=True)
                seen += 1

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "\u0410. \u0412. \u0411\u0435\u0440\u0435\u0437\u0443\u0442\u0441\u043a\u0438\u0439" in text:
            set_text(paragraph, text.replace("\u0410. \u0412. \u0411\u0435\u0440\u0435\u0437\u0443\u0442\u0441\u043a\u0438\u0439", "\u0410. \u0412. \u0421\u0438\u0434\u043e\u0440\u043e\u0432"))
        elif text == "\u041a\u041b\u0418\u0415\u041d\u0422\u0421\u041a\u0410\u042f \u0418 \u0421\u0415\u0420\u0412\u0415\u0420\u041d\u0410\u042f \u0427\u0410\u0421\u0422\u042c":
            set_text(paragraph, "\u041c\u041e\u0414\u0423\u041b\u042c \u0410\u041d\u0410\u041b\u0418\u0417\u0410 \u041e\u0422\u0412\u0415\u0422\u041e\u0412 \u041f\u041e\u041b\u042c\u0417\u041e\u0412\u0410\u0422\u0415\u041b\u042f", center=True)

    # Add missing fourth title line after the module title in the second title block if not present nearby.
    for i, paragraph in enumerate(list(doc.paragraphs)):
        if paragraph.text.strip() == "\u041c\u041e\u0414\u0423\u041b\u042c \u0410\u041d\u0410\u041b\u0418\u0417\u0410 \u041e\u0422\u0412\u0415\u0422\u041e\u0412 \u041f\u041e\u041b\u042c\u0417\u041e\u0412\u0410\u0422\u0415\u041b\u042f":
            nxt = doc.paragraphs[i + 1].text.strip() if i + 1 < len(doc.paragraphs) else ""
            if nxt != MODULE_TITLE_LINES[3]:
                added = insert_paragraph_after(paragraph, MODULE_TITLE_LINES[3])
                normalize_paragraph(added, center=True)


CONTENT: list[tuple[str, str]] = [
    ("normal_center", "\u0413\u041b\u041e\u0421\u0421\u0410\u0420\u0418\u0419"),
    ("normal", "API (Application Programming Interface) - программный интерфейс приложения, набор методов для взаимодействия между компонентами."),
    ("normal", "Backend-приложение - основная серверная часть Interview Coach, которая хранит вопросы, проводит сессию интервью и передает завершенные вопросы и ответы в модуль анализа."),
    ("normal", "FastAPI - web-фреймворк Python, используемый для реализации REST API модуля анализа ответов."),
    ("normal", "HFProvider - компонент модуля, который локально загружает Qwen/Qwen2.5-3B-Instruct и LoRA/QLoRA-адаптер через Transformers и PEFT."),
    ("normal", "LLM - большая языковая модель, применяемая для анализа технических ответов и формирования структурированного отчета."),
    ("normal", "LoRA/QLoRA - методы параметро-эффективного дообучения модели, используемые для адаптации базовой LLM к задаче оценки ответов интервью."),
    ("normal", "OpenAPI - машиночитаемое описание REST API, используемое для интеграции модуля анализа с основным backend-приложением."),
    ("normal", "PostgreSQL - система управления базами данных, используемая для хранения статусов задач и сформированных отчетов."),
    ("normal", "Сессия собеседования - завершенный набор вопросов и ответов пользователя в рамках выбранной специализации и грейда."),
    ("normal", "Рубрика оценки - набор критериев и весов, по которым оцениваются корректность, полнота, ясность, практичность и терминология ответа."),
    ("normal", "Job - задача анализа, создаваемая при обращении к endpoint формирования отчета."),
    ("normal", "MAE, RMSE, bias, within_10, within_15 - метрики качества модели, используемые для сравнения оценок модели с экспертной разметкой."),
    ("blank", ""),
    ("normal_center", "\u0410\u041d\u041d\u041e\u0422\u0410\u0426\u0418\u042f"),
    ("normal", "Руководство оператора предназначено для организации корректной эксплуатации модуля анализа ответов пользователя с использованием большой языковой модели, входящего в состав веб-приложения Interview Coach."),
    ("normal", "Модуль принимает от основного backend-приложения завершенную сессию собеседования, включающую тексты вопросов и ответы пользователя, выполняет анализ через локальную LLM и возвращает структурированный JSON-отчет с оценками, сильными сторонами, проблемами и рекомендациями."),
    ("normal", "В руководстве описаны назначение программы, условия выполнения, порядок запуска через Docker и локальное окружение, работа через REST API и консоль, проверка качества модели, остановка сервиса, просмотр журналов и типовые сообщения оператору."),
    ("blank", ""),
    ("normal_center", "\u0421\u041e\u0414\u0415\u0420\u0416\u0410\u041d\u0418\u0415"),
    ("normal", "НАЗНАЧЕНИЕ ПРОГРАММЫ"),
    ("normal", "УСЛОВИЯ ВЫПОЛНЕНИЯ ПРОГРАММЫ"),
    ("normal", "ВЫПОЛНЕНИЕ ПРОГРАММЫ"),
    ("normal", "СООБЩЕНИЯ ОПЕРАТОРУ"),
    ("normal", "ПРИЛОЖЕНИЕ 1. СПИСОК ИСПОЛЬЗУЕМОЙ ЛИТЕРАТУРЫ"),
    ("blank", ""),
    ("h1", "\u041d\u0410\u0417\u041d\u0410\u0427\u0415\u041d\u0418\u0415 \u041f\u0420\u041e\u0413\u0420\u0410\u041c\u041c\u042b"),
    ("h2", "Функциональное назначение"),
    ("normal", "Модуль анализа ответов предназначен для автоматизированной оценки ответов пользователя после прохождения технического собеседования. Он не проводит интервью самостоятельно и не является основным хранилищем вопросов в production-контуре."),
    ("normal", "Основное backend-приложение хранит вопросы, управляет сессией и отправляет в модуль анализа итоговый JSON с полями requestId, sessionId, clientId, scenario и items. Каждый элемент items содержит itemId, questionId, questionText, answerText, askedAt и tags."),
    ("normal", "Модуль формирует отчет, включающий общий балл, оценки по критериям, разбор каждого вопроса, тематические выводы, рекомендации, версии модели, промпта, рубрик и источника вопросов."),
    ("normal", "Если questionId отсутствует в локальном учебном банке, модуль не завершает работу ошибкой. В этом случае используется runtime-карточка и runtime-рубрика, сформированные на основе текста вопроса, тегов и тем сценария."),
    ("h2", "Эксплуатационное назначение"),
    ("normal", "Программа используется оператором для запуска и сопровождения API-сервиса анализа ответов, проверки его готовности, передачи тестовых сессий, просмотра журналов, диагностики ошибок и контроля качества модели."),
    ("normal", "Для демонстрации и проверки без основного backend-приложения предусмотрен консольный режим. Он позволяет вывести вопросы из локального учебного набора, пройти интервью вручную, сформировать JSON-отчет и запустить оценку качества на тестовой выборке."),
    ("normal", "Локальные карточки вопросов и датасет применяются для обучения, демонстрации и проверки качества. В producton-интеграции источник вопросов находится на стороне backend-приложения."),
    ("blank", ""),
    ("h1", "\u0423\u0421\u041b\u041e\u0412\u0418\u042f \u0412\u042b\u041f\u041e\u041b\u041d\u0415\u041d\u0418\u042f \u041f\u0420\u041e\u0413\u0420\u0410\u041c\u041c\u042b"),
    ("h2", "Минимальный состав аппаратурных средств"),
    ("normal", "Для запуска API без локального inference достаточно сервера или рабочей станции с 2 vCPU, 4 ГБ оперативной памяти и 20 ГБ свободного дискового пространства."),
    ("normal", "Для полного запуска HFProvider с моделью Qwen/Qwen2.5-3B-Instruct и LoRA/QLoRA-адаптером рекомендуется рабочая станция или сервер с 8 vCPU, 24 ГБ оперативной памяти и 100 ГБ дискового пространства."),
    ("normal", "При наличии NVIDIA GPU рекомендуется использовать видеокарту с поддержкой CUDA. GPU-режим ускоряет генерацию отчета и снижает задержку обработки сессии."),
    ("h2", "Минимальный состав программных средств"),
    ("normal", "Для Docker-развертывания необходимы Docker Engine или Docker Desktop, Docker Compose, доступ к образам Python и PostgreSQL, а также сетевой доступ для загрузки базовой модели при первом HF-запуске."),
    ("normal", "Для локального запуска без Docker необходимы Python 3.12+, виртуальное окружение, FastAPI, Uvicorn, Pydantic, psycopg, PyTorch, Transformers, PEFT, bitsandbytes и зависимости проекта, устанавливаемые из pyproject.toml."),
    ("normal", "Для хранения статусов задач и отчетов используется PostgreSQL. Для локальной проверки допускается in-memory-хранилище, однако для интеграционного и серверного контура рекомендуется PostgreSQL."),
    ("h2", "Требования к квалификации персонала"),
    ("normal", "Оператор должен владеть базовыми навыками работы с командной строкой, Docker, HTTP-запросами, переменными окружения и просмотром журналов контейнеров."),
    ("normal", "Для диагностики LLM-компонента оператор должен понимать назначение режимов mock, ollama и hf, знать путь к адаптеру модели и уметь проверять доступность endpoint health."),
    ("normal", "Для сопровождения серверной интеграции оператор должен иметь доступ к репозиторию, файлу .env, значению X-API-Key, параметрам PostgreSQL и журналам основного backend-приложения."),
    ("blank", ""),
    ("h1", "\u0412\u042b\u041f\u041e\u041b\u041d\u0415\u041d\u0418\u0415 \u041f\u0420\u041e\u0413\u0420\u0410\u041c\u041c\u042b"),
    ("h2", "Подготовка к работе"),
    ("normal", "Перед запуском оператор клонирует репозиторий модуля анализа, переходит в каталог проекта, создает файл .env на основе .env.example и задает значения переменных окружения."),
    ("normal", "Оператор проверяет наличие каталога training/artifacts/qwen2.5-3b-interview-full-ru-qlora-v1, свободный порт 8000, доступность PostgreSQL и корректность API-ключа для заголовка X-API-Key."),
    ("normal", "Минимальная последовательность команд подготовки:"),
    ("normal", "git clone <url_репозитория>"),
    ("normal", "cd analysis_module"),
    ("normal", "copy .env.example .env"),
    ("h2", "Локальный запуск через Docker"),
    ("normal", "Основной способ запуска итоговой версии модуля - Docker Compose. В этом режиме поднимаются сервис анализа ответов и PostgreSQL."),
    ("normal", "Для запуска в Windows используется команда:"),
    ("normal", ".\\start.bat"),
    ("normal", "Альтернативная команда без batch-файла:"),
    ("normal", "docker compose up -d --build"),
    ("normal", "Для запуска на сервере с NVIDIA GPU используется команда:"),
    ("normal", "docker compose -f docker-compose.yml -f docker-compose.gpu.yml up -d --build"),
    ("normal", "После запуска оператор проверяет состояние контейнеров и доступность API:"),
    ("normal", "docker compose ps"),
    ("normal", "curl http://127.0.0.1:8000/assessment/v1/health"),
    ("h2", "Локальный запуск HF API без Docker"),
    ("normal", "При запуске без Docker оператор создает виртуальное окружение и устанавливает зависимости проекта."),
    ("normal", "python -m venv .venv"),
    ("normal", ".\\.venv\\Scripts\\activate"),
    ("normal", "pip install -e \".[dev,hf_runtime]\""),
    ("normal", "Для старта API в HF/LoRA-режиме используется команда:"),
    ("normal", ".\\start_hf_api.bat"),
    ("normal", "После запуска доступны Swagger UI по адресу http://127.0.0.1:8000/docs, OpenAPI JSON по адресу http://127.0.0.1:8000/openapi.json и endpoint состояния /assessment/v1/health."),
    ("h2", "Работа через REST API"),
    ("normal", "REST API используется основным backend-приложением. Все защищенные запросы должны содержать заголовок X-API-Key."),
    ("normal", "Основной endpoint анализа: POST /assessment/v1/report. Входной JSON должен содержать requestId, sessionId, clientId, mode, scenario и массив items."),
    ("normal", "В синхронном режиме mode=sync ответ содержит status=ready, объект job и объект report. В асинхронном режиме mode=async первичный ответ возвращается со статусом 202, после чего готовность проверяется через GET /assessment/v1/report/{jobId}/status."),
    ("normal", "Пример минимального запроса:"),
    ("normal", "curl -X POST \"http://127.0.0.1:8000/assessment/v1/report\" -H \"X-API-Key: demo-api-key\" -H \"Content-Type: application/json\" -d \"{\\\"requestId\\\":\\\"req-1001\\\",\\\"sessionId\\\":\\\"session-1001\\\",\\\"clientId\\\":\\\"main-backend\\\",\\\"mode\\\":\\\"sync\\\",\\\"scenario\\\":{\\\"specialization\\\":\\\"backend\\\",\\\"grade\\\":\\\"junior\\\",\\\"reportLanguage\\\":\\\"ru\\\"},\\\"items\\\":[{\\\"itemId\\\":\\\"item-1\\\",\\\"questionId\\\":\\\"q-1\\\",\\\"questionText\\\":\\\"Что такое REST?\\\",\\\"answerText\\\":\\\"REST - архитектурный стиль для проектирования HTTP API.\\\",\\\"tags\\\":[\\\"http\\\"]}]}\""),
    ("h2", "Работа через консоль"),
    ("normal", "Консольный режим используется для демонстрации, локальной проверки и подготовки материалов без графического интерфейса."),
    ("normal", "Получение списка вопросов из локального учебного набора:"),
    ("normal", ".\\console.bat questions --specialization backend --grade junior --limit 10"),
    ("normal", "Быстрая проверка пайплайна без загрузки LLM:"),
    ("normal", ".\\console.bat sample --quality good --specialization backend --grade junior --limit 10 --llm mock --output training\\reports\\mock_report.json"),
    ("normal", "Ручное прохождение интервью через локальную модель:"),
    ("normal", ".\\console.bat run --specialization backend --grade junior --limit 10 --llm hf --output training\\reports\\manual_hf_report.json"),
    ("h2", "Проверка качества модели"),
    ("normal", "Качество модели проверяется на размеченной тестовой выборке. Оператор запускает генерацию предсказаний и расчет метрик."),
    ("normal", ".\\console.bat evaluate --limit 5 --llm hf --output training\\reports\\hf_eval_predictions.jsonl"),
    ("normal", ".\\.venv\\Scripts\\python.exe training\\scripts\\evaluate_predictions.py --input training\\reports\\hf_eval_predictions.jsonl --output training\\reports\\hf_quality_metrics.json"),
    ("normal", "Основные метрики: MAE - средняя абсолютная ошибка, RMSE - среднеквадратичная ошибка, bias - среднее смещение оценки, within_10 и within_15 - доля ответов, где оценка модели отличается от экспертной не более чем на 10 или 15 баллов."),
    ("h2", "Остановка, обновление и диагностика"),
    ("normal", "Для остановки Docker-режима используется команда:"),
    ("normal", ".\\stop.bat"),
    ("normal", "Для просмотра журналов используются команды:"),
    ("normal", "docker compose logs -f analysis-module"),
    ("normal", "docker compose logs --tail 200 analysis-module"),
    ("normal", "docker compose logs --tail 200 postgres"),
    ("normal", "При обновлении LoRA/QLoRA-адаптера оператор проверяет переменную ANALYSIS_HF_ADAPTER_PATH, после чего пересоздает контейнер analysis-module."),
    ("normal", "docker compose up -d --build --force-recreate analysis-module"),
    ("blank", ""),
    ("h1", "\u0421\u041e\u041e\u0411\u0429\u0415\u041d\u0418\u042f \u041e\u041f\u0415\u0420\u0410\u0422\u041e\u0420\u0423"),
    ("normal", "При ошибках модуль возвращает структурированный JSON-ответ с полями status, errorCode, message и details. По этим полям оператор определяет причину сбоя и дальнейшие действия."),
    ("normal", "ok - сервис доступен, health-запрос выполнен успешно."),
    ("normal", "ready - отчет сформирован и доступен в поле report."),
    ("normal", "created - задача зарегистрирована и ожидает обработки."),
    ("normal", "processing - задача анализа выполняется."),
    ("normal", "AUTHENTICATION_ERROR - отсутствует или неверно указан X-API-Key. Нужно проверить ANALYSIS_API_KEY и настройки вызывающего backend-приложения."),
    ("normal", "INVALID_INPUT - запрос не прошел валидацию. Типичные причины: отсутствуют requestId, sessionId, clientId, questionText или answerText; превышено ограничение 20 вопросов; ответ длиннее 4000 символов."),
    ("normal", "REQUEST_CONFLICT - один и тот же requestId повторно используется с другим телом запроса. Нужно передать новый requestId или повторить исходный запрос без изменений."),
    ("normal", "JOB_NOT_FOUND - задача с указанным jobId отсутствует в хранилище. Нужно проверить корректность jobId и состояние PostgreSQL."),
    ("normal", "REPORT_NOT_READY - отчет еще формируется. Нужно повторить запрос статуса позднее."),
    ("normal", "INVALID_MODEL_OUTPUT - модель вернула невалидный JSON. Нужно проверить журналы analysis-module, параметры генерации и доступность HFProvider."),
    ("normal", "MODEL_TIMEOUT - модель превысила допустимое время генерации. Нужно проверить загрузку сервера, режим hf, наличие GPU и длину входной сессии."),
    ("normal", "DATABASE_UNAVAILABLE или DATABASE_WRITE_ERROR - PostgreSQL недоступен либо запись статуса задачи невозможна. Нужно проверить контейнер postgres и ANALYSIS_DATABASE_URL."),
    ("normal", "INTERNAL - непредвиденная ошибка. Нужно сохранить requestId, jobId, время ошибки и фрагмент журнала для передачи разработчику."),
    ("blank", ""),
    ("h1", "\u041f\u0420\u0418\u041b\u041e\u0416\u0415\u041d\u0418\u0415 1"),
    ("normal", "\u0421\u041f\u0418\u0421\u041e\u041a \u0418\u0421\u041f\u041e\u041b\u042c\u0417\u0423\u0415\u041c\u041e\u0419 \u041b\u0418\u0422\u0415\u0420\u0410\u0422\u0423\u0420\u042b"),
    ("normal", "ГОСТ 19.101-77. Виды программ и программных документов. ЕСПД."),
    ("normal", "ГОСТ 19.102-77. Стадии разработки. ЕСПД."),
    ("normal", "ГОСТ 19.103-77. Обозначения программ и программных документов. ЕСПД."),
    ("normal", "ГОСТ 19.104-78. Основные надписи. ЕСПД."),
    ("normal", "ГОСТ 19.105-78. Общие требования к программным документам. ЕСПД."),
    ("normal", "ГОСТ 19.106-78. Требования к программным документам, выполненным печатным способом. ЕСПД."),
    ("normal", "ГОСТ 19.505-79. Руководство оператора. Требования к содержанию и оформлению. ЕСПД."),
    ("normal", "Python Documentation. Электронный ресурс."),
    ("normal", "FastAPI Documentation. Электронный ресурс."),
    ("normal", "Pydantic Documentation. Электронный ресурс."),
    ("normal", "PostgreSQL Documentation. Электронный ресурс."),
    ("normal", "Docker Documentation. Электронный ресурс."),
    ("normal", "OpenAPI Specification. Электронный ресурс."),
    ("normal", "Hugging Face Transformers Documentation. Электронный ресурс."),
    ("normal", "Hugging Face PEFT Documentation. Электронный ресурс."),
    ("normal", "PyTorch Documentation. Электронный ресурс."),
    ("normal", "Qwen/Qwen2.5-3B-Instruct model card. Электронный ресурс."),
]


def replace_body(doc: Document) -> None:
    start = find_paragraph(doc, "\u0413\u041b\u041e\u0421\u0421\u0410\u0420\u0418\u0419")
    end = find_paragraph(doc, "\u041b\u0418\u0421\u0422 \u0420\u0415\u0413\u0418\u0421\u0422\u0420\u0410\u0426\u0418\u0418 \u0418\u0417\u041c\u0415\u041d\u0415\u041d\u0418\u0419")
    start_idx = next(i for i, paragraph in enumerate(doc.paragraphs) if paragraph._p is start._p)
    end_idx = next(i for i, paragraph in enumerate(doc.paragraphs) if paragraph._p is end._p)
    anchor = doc.paragraphs[start_idx - 1]
    for paragraph in list(doc.paragraphs[start_idx:end_idx]):
        remove_paragraph(paragraph)

    current = anchor
    for kind, text in CONTENT:
        if kind == "blank":
            current = insert_paragraph_after(current, "")
        elif kind == "h1":
            current = insert_paragraph_after(current, text, "Heading 1")
        elif kind == "h2":
            current = insert_paragraph_after(current, text, "Heading 2")
        elif kind == "normal_center":
            current = insert_paragraph_after(current, text)
            normalize_paragraph(current, center=True)
        else:
            current = insert_paragraph_after(current, text)


def normalize_dashes(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace("—", "-").replace("–", "-")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace("—", "-").replace("–", "-")


def main() -> None:
    if TARGET.exists() and not BACKUP.exists():
        shutil.copy2(TARGET, BACKUP)
    shutil.copy2(REFERENCE, TARGET)

    doc = Document(TARGET)
    replace_title_pages(doc)
    replace_body(doc)
    normalize_dashes(doc)
    doc.save(TARGET)
    print(f"target={TARGET}")
    print(f"backup={BACKUP}")


if __name__ == "__main__":
    main()
