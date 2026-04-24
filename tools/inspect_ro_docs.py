from pathlib import Path

from docx import Document


BASE = Path("C:/Users/super")
REFERENCE = next((BASE / "Downloads/Telegram Desktop").glob("РО Березутск*.docx"))
CURRENT = (
    BASE
    / "Desktop"
    / "\u0414\u0438\u043f\u043b\u043e\u043c"
    / "\u0444\u0430\u0439\u043d\u0430\u043b\u0438 \u0434\u043e\u043a\u0438"
    / "\u0420\u041e_\u0421\u0438\u0434\u043e\u0440\u043e\u0432.docx"
)
OUT = Path("tools/ro_docs_outline.txt")


def clean(text: str) -> str:
    return " ".join(text.split())


def dump(path: Path, title: str) -> list[str]:
    doc = Document(path)
    lines = [f"==== {title}: {path}", f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}"]
    for i, paragraph in enumerate(doc.paragraphs):
        text = clean(paragraph.text)
        if text:
            lines.append(f"P{i:04d} [{paragraph.style.name}]: {text[:350]}")
    for ti, table in enumerate(doc.tables):
        lines.append(f"TABLE {ti}: rows={len(table.rows)} cols={len(table.columns)}")
        for ri, row in enumerate(table.rows[:12]):
            cells = [clean(cell.text)[:140] for cell in row.cells]
            lines.append(f"  R{ri:02d}: " + " | ".join(cells))
    full = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full += "\n" + cell.text
    for token in ["—", "–", "SQLAlchemy", "snake_case", "request_id", "question_id", "UNKNOWN_QUESTION", "Ollama", "ngrok"]:
        lines.append(f"COUNT {token!r}: {full.count(token)}")
    return lines


OUT.write_text("\n".join(dump(REFERENCE, "REFERENCE") + [""] + dump(CURRENT, "CURRENT")), encoding="utf-8")
print(OUT.resolve())
