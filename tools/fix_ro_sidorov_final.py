from pathlib import Path

from docx import Document
from docx.shared import Pt


TARGET = (
    Path("C:/Users/super/Desktop")
    / "\u0414\u0438\u043f\u043b\u043e\u043c"
    / "\u0444\u0430\u0439\u043d\u0430\u043b\u0438 \u0434\u043e\u043a\u0438"
    / "\u0420\u041e_\u0421\u0438\u0434\u043e\u0440\u043e\u0432.docx"
)

TITLE_LINES = [
    "\u0412\u0415\u0411-\u041f\u0420\u0418\u041b\u041e\u0416\u0415\u041d\u0418\u0415 \u0414\u041b\u042f \u0418\u041c\u0418\u0422\u0410\u0426\u0418\u0418 \u0421\u041e\u0411\u0415\u0421\u0415\u0414\u041e\u0412\u0410\u041d\u0418\u0419",
    "\u0421 \u0418\u041d\u0422\u0415\u041b\u041b\u0415\u041a\u0422\u0423\u0410\u041b\u042c\u041d\u041e\u0419 \u041e\u0411\u0420\u0410\u0422\u041d\u041e\u0419 \u0421\u0412\u042f\u0417\u042c\u042e:",
    "\u041c\u041e\u0414\u0423\u041b\u042c \u0410\u041d\u0410\u041b\u0418\u0417\u0410 \u041e\u0422\u0412\u0415\u0422\u041e\u0412 \u041f\u041e\u041b\u042c\u0417\u041e\u0412\u0410\u0422\u0415\u041b\u042f",
    "\u0421 \u0418\u0421\u041f\u041e\u041b\u042c\u0417\u041e\u0412\u0410\u041d\u0418\u0415\u041c \u0411\u041e\u041b\u042c\u0428\u041e\u0419 \u042f\u0417\u042b\u041a\u041e\u0412\u041e\u0419 \u041c\u041e\u0414\u0415\u041b\u0418",
]

REPLACEMENTS = {
    "\u0410. \u0412. \u0411\u0435\u0440\u0435\u0437\u0443\u0442\u0441\u043a\u0438\u0439": "\u0410. \u0412. \u0421\u0438\u0434\u043e\u0440\u043e\u0432",
    "\u041a\u041b\u0418\u0415\u041d\u0422\u0421\u041a\u0410\u042f \u0418 \u0421\u0415\u0420\u0412\u0415\u0420\u041d\u0410\u042f \u0427\u0410\u0421\u0422\u042c": "\u041c\u041e\u0414\u0423\u041b\u042c \u0410\u041d\u0410\u041b\u0418\u0417\u0410 \u041e\u0422\u0412\u0415\u0422\u041e\u0412 \u041f\u041e\u041b\u042c\u0417\u041e\u0412\u0410\u0422\u0415\u041b\u042f",
    "\u041a\u043b\u0438\u0435\u043d\u0442\u0441\u043a\u0430\u044f \u0438 \u0441\u0435\u0440\u0432\u0435\u0440\u043d\u0430\u044f \u0447\u0430\u0441\u0442\u044c": "\u041c\u043e\u0434\u0443\u043b\u044c \u0430\u043d\u0430\u043b\u0438\u0437\u0430 \u043e\u0442\u0432\u0435\u0442\u043e\u0432",
    "\u041a\u041b\u0418\u0415\u041d\u0422\u0421\u041a\u0410\u042f": "\u041c\u041e\u0414\u0423\u041b\u042c",
}


def set_font(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def replace_text_in_paragraph(paragraph) -> None:
    if "?" in paragraph.text and 55 <= doc.paragraphs.index(paragraph) <= 59:
        return
    text = paragraph.text
    for old, new in REPLACEMENTS.items():
        text = text.replace(old, new)
    text = text.replace(chr(8212), "-").replace(chr(8211), "-")
    if text != paragraph.text:
        paragraph.text = text
        set_font(paragraph)
    else:
        for run in paragraph.runs:
            run.text = run.text.replace(chr(8212), "-").replace(chr(8211), "-")


doc = Document(TARGET)

for idx, text in zip([27, 28, 29, 30], TITLE_LINES, strict=True):
    paragraph = doc.paragraphs[idx]
    paragraph.text = text
    set_font(paragraph)

for idx, text in zip([56, 57, 58, 59], TITLE_LINES, strict=True):
    paragraph = doc.paragraphs[idx]
    paragraph.text = text
    set_font(paragraph)

for paragraph in doc.paragraphs:
    replace_text_in_paragraph(paragraph)
    if paragraph.text.strip().startswith("\u041b\u0438\u0441\u0442\u043e\u0432:"):
        paragraph.text = "\u041b\u0438\u0441\u0442\u043e\u0432: 22"
        set_font(paragraph)

for section in doc.sections:
    for container in (section.header, section.footer):
        for paragraph in container.paragraphs:
            replace_text_in_paragraph(paragraph)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph)

doc.save(TARGET)
print(TARGET)
